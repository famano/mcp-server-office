import os
import pytest
from mcp_server_office.office import validate_path, read_docx, write_docx, edit_docx_paragraph as edit_docx, edit_docx_insert, extract_table_text
from docx import Document
from docx.table import Table
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

@pytest.fixture
def sample_docx_with_track_changes():
    """Create a sample docx file with track changes for testing."""
    path = "test_track_changes.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    
    # Add text with track changes
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = "Original"
    run.append(text)
    paragraph._element.append(run)
    
    # Add deletion
    del_element = OxmlElement('w:del')
    del_element.set(qn('w:author'), 'Test Author')
    del_element.set(qn('w:date'), '2024-01-27T00:00:00Z')
    del_run = OxmlElement('w:r')
    del_text = OxmlElement('w:delText')
    del_text.text = " deleted"
    del_run.append(del_text)
    del_element.append(del_run)
    paragraph._element.append(del_element)
    
    # Add insertion
    ins_element = OxmlElement('w:ins')
    ins_element.set(qn('w:author'), 'Test Author')
    ins_element.set(qn('w:date'), '2024-01-27T00:00:00Z')
    ins_run = OxmlElement('w:r')
    ins_text = OxmlElement('w:t')
    ins_text.text = " inserted"
    ins_run.append(ins_text)
    ins_element.append(ins_run)
    paragraph._element.append(ins_element)
    
    doc.save(path)
    yield path
    if os.path.exists(path):
        os.remove(path)

@pytest.fixture
def sample_docx():
    """Create a sample docx file for testing."""
    path = "test_sample.docx"
    doc = Document()
    doc.add_paragraph("Hello World")
    
    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    
    doc.add_paragraph("Goodbye World")
    doc.save(path)
    yield path
    # Cleanup
    if os.path.exists(path):
        os.remove(path)

@pytest.mark.asyncio
async def test_validate_path(sample_docx):
    """Test path validation."""
    assert await validate_path(os.path.abspath(sample_docx)) == True
    with pytest.raises(ValueError):
        await validate_path(sample_docx) 
    with pytest.raises(ValueError):
        await validate_path("nonexistent.docx")

@pytest.mark.asyncio
async def test_read_docx_with_track_changes(sample_docx_with_track_changes):
    """Test reading docx file with track changes."""
    content = await read_docx(os.path.abspath(sample_docx_with_track_changes))
    assert "Original" in content
    assert not "deleted" in content
    assert "inserted" in content

@pytest.mark.asyncio
async def test_read_docx(sample_docx):
    """Test reading docx file."""
    content = await read_docx(os.path.abspath(sample_docx))
    assert "Hello World" in content
    assert "Goodbye World" in content
    assert "[Table]" in content
    assert "A1 | B1" in content
    assert "A2 | B2" in content

    with pytest.raises(ValueError):
        await read_docx("nonexistent.docx")

@pytest.mark.asyncio
async def test_write_docx():
    """Test writing docx file."""
    test_path = "test_write.docx"
    test_content = "Test Paragraph\n\n[Table]\nC1 | D1\nC2 | D2\n\nFinal Paragraph"
    
    try:
        await write_docx(test_path, test_content)
        assert os.path.exists(test_path)
        
        # Verify content
        doc = Document(test_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        assert "Test Paragraph" in paragraphs
        assert "Final Paragraph" in paragraphs
        
        # Verify table
        table = doc.tables[0]
        assert table.cell(0, 0).text == "C1"
        assert table.cell(0, 1).text == "D1"
        assert table.cell(1, 0).text == "C2"
        assert table.cell(1, 1).text == "D2"
    
    finally:
        if os.path.exists(test_path):
            os.remove(test_path)

@pytest.mark.asyncio
async def test_edit_docx_with_track_changes(sample_docx_with_track_changes):
    """Test editing docx file with track changes."""
    abs_path = os.path.abspath(sample_docx_with_track_changes)
    await edit_docx(abs_path, [{"paragraph_index": 0,"search": "Original", "replace": "Modified"}])
    
    # Verify track changes are preserved
    content = await read_docx(abs_path)
    assert not "Original" in content
    assert "Modified" in content
    assert not "deleted" in content
    assert "inserted" in content

@pytest.fixture
def complex_docx():
    """Create a sample docx file with complex content for testing cross-paragraph and table edits."""
    path = "test_complex.docx"
    doc = Document()
    
    # Add paragraphs with text that spans multiple paragraphs
    doc.add_paragraph("First part of")
    doc.add_paragraph("a sentence that")
    doc.add_paragraph("spans multiple paragraphs")
    
    doc.add_paragraph("Some text before table")
    
    # Add table with text that will be edited
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Table"
    table.cell(0, 1).text = "Content"
    table.cell(1, 0).text = "More"
    table.cell(1, 1).text = "Text"
    
    doc.add_paragraph("Some text after table")
    
    doc.save(path)
    yield path
    if os.path.exists(path):
        os.remove(path)

@pytest.fixture
def formatted_docx():
    """Create a sample docx file with formatted text for testing."""
    path = "test_formatted.docx"
    doc = Document()
    
    # Add paragraph with formatted text
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Bold")
    run.bold = True
    run = paragraph.add_run(" and ")
    run = paragraph.add_run("Italic")
    run.italic = True
    run = paragraph.add_run(" text")
    
    doc.save(path)
    yield path
    if os.path.exists(path):
        os.remove(path)

@pytest.mark.asyncio
async def test_edit_docx(sample_docx):
    abs_sample_docx = os.path.abspath(sample_docx)
    
    """Test editing docx file with paragraph index."""
    # Test single edit with paragraph index
    result = await edit_docx(abs_sample_docx, [{"paragraph_index": 0, "search": "Hello", "replace": "Hi"}])
    assert "-Hello" in result
    assert "+Hi" in result
    
    # Test multiple edits in different paragraphs
    result = await edit_docx(abs_sample_docx, [
        {"paragraph_index": 0, "search": "Hi", "replace": "Hellow"},
        {"paragraph_index": 2, "search": "Goodbye", "replace": "Bye"}
    ])
    assert "-Hi" in result
    assert "+Hellow" in result
    assert "-Goodbye" in result
    assert "+Bye" in result
    
    # Test non-existent text in paragraph
    with pytest.raises(ValueError) as exc_info:
        await edit_docx(abs_sample_docx, [{"paragraph_index": 0, "search": "NonexistentText", "replace": "NewText"}])
    assert "'NonexistentText' in paragraph 0" in str(exc_info.value)
    
    # Test invalid paragraph index
    with pytest.raises(ValueError) as exc_info:
        await edit_docx(abs_sample_docx, [{"paragraph_index": 999, "search": "Hello", "replace": "Hi"}])
    assert "Paragraph index out of range: 999" in str(exc_info.value)
    
    # Test invalid file
    with pytest.raises(ValueError):
        await edit_docx("nonexistent.docx", [{"paragraph_index": 0, "search": "Hello", "replace": "Hi"}])

@pytest.mark.asyncio
async def test_edit_docx_format_preservation(formatted_docx):
    """Test that formatting is preserved when editing text."""
    abs_formatted_docx = os.path.abspath(formatted_docx)
    
    # Edit text while preserving bold formatting
    result = await edit_docx(abs_formatted_docx, [
        {"paragraph_index": 0, "search": "Bold and Italic text", "replace": "Modified text"}
    ])
    
    # Verify content was changed
    assert "-Bold and Italic text" in result
    assert "+Modified text" in result
    
    # Verify formatting was preserved by checking the document directly
    doc = Document(abs_formatted_docx)
    paragraph = doc.paragraphs[0]
    assert len(paragraph.runs) == 1  # Should be consolidated into a single run
    run = paragraph.runs[0]
    assert run.bold  # Should inherit bold formatting from first run

@pytest.mark.asyncio
async def test_edit_docx_table_content(complex_docx):
    """Test editing text within table cells."""
    abs_complex_docx = os.path.abspath(complex_docx)
    
    # Test editing table content
    result = await edit_docx(abs_complex_docx, [
        {"paragraph_index": 4,"search": "Table | Content", "replace": "Modified | Cell"},
        {"paragraph_index": 4,"search": "More | Text", "replace": "Modification | Two"},
    ])
    assert "-Table | Content" in result
    assert "+Modified | Cell" in result
    assert "-More | Text" in result
    assert "+Modification | Two" in result


@pytest.fixture
def table_at_start_docx():
    """Create a docx file that starts with a table."""
    path = "test_table_at_start.docx"
    doc = Document()
    
    # Add table at the start
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Start1"
    table.cell(0, 1).text = "Start2"
    table.cell(1, 0).text = "Start3"
    table.cell(1, 1).text = "Start4"
    
    # Add some text after table
    doc.add_paragraph("Text after table")
    
    doc.save(path)
    yield path
    if os.path.exists(path):
        os.remove(path)

@pytest.fixture
def table_after_empty_paragraph_docx():
    """Create a docx file with empty paragraph before table."""
    path = "test_table_after_empty.docx"
    doc = Document()
    
    # Add empty paragraph
    doc.add_paragraph("")
    
    # Add table after empty paragraph
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Empty1"
    table.cell(0, 1).text = "Empty2"
    table.cell(1, 0).text = "Empty3"
    table.cell(1, 1).text = "Empty4"
    
    doc.save(path)
    yield path
    if os.path.exists(path):
        os.remove(path)

@pytest.mark.asyncio
async def test_edit_docx_table_at_start(table_at_start_docx):
    """Test editing table that appears at the start of document."""
    abs_path = os.path.abspath(table_at_start_docx)
    
    # Test editing the table content
    result = await edit_docx(abs_path, [
        {"paragraph_index": 0, "search": "Start1 | Start2", "replace": "Modified1 | Modified2"}
    ])
    
    # Verify changes
    assert "-Start1 | Start2" in result
    assert "+Modified1 | Modified2" in result
    
    # Verify the document structure is maintained
    content = await read_docx(abs_path)
    assert "Text after table" in content  # Verify text after table is preserved

@pytest.mark.asyncio
async def test_edit_docx_table_after_empty_paragraph(table_after_empty_paragraph_docx):
    """Test editing table that appears after an empty paragraph."""
    abs_path = os.path.abspath(table_after_empty_paragraph_docx)
    
    # Test editing the table content
    result = await edit_docx(abs_path, [
        {"paragraph_index": 1, "search": "Empty1 | Empty2", "replace": "Modified1 | Modified2"}
    ])
    
    # Verify changes
    assert "-Empty1 | Empty2" in result
    assert "+Modified1 | Modified2" in result
    
    # Verify the document structure is maintained
    content = await read_docx(abs_path)
    assert "Modified1" in content
    assert "Modified2" in content

@pytest.fixture
def deleted_text_before_table_docx():
    """Create a docx file with deleted text in a paragraph before table."""
    path = "test_deleted_text_before_table.docx"
    doc = Document()
    
    # 削除された文字を含む段落を追加
    paragraph = doc.add_paragraph()
    # 削除マークを付けたテキストを追加
    del_element = OxmlElement('w:del')
    del_element.set(qn('w:author'), 'Test Author')
    del_element.set(qn('w:date'), '2024-01-27T00:00:00Z')
    del_run = OxmlElement('w:r')
    del_text = OxmlElement('w:delText')
    del_text.text = "This text is deleted"
    del_run.append(del_text)
    del_element.append(del_run)
    paragraph._element.append(del_element)
    
    # テーブルを追加
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Test1"
    table.cell(0, 1).text = "Test2"
    table.cell(1, 0).text = "Test3"
    table.cell(1, 1).text = "Test4"
    
    doc.save(path)
    yield path
    if os.path.exists(path):
        os.remove(path)

@pytest.mark.asyncio
async def test_edit_docx_table_after_deleted_text(deleted_text_before_table_docx):
    """Test editing table that appears after a paragraph with deleted text."""
    abs_path = os.path.abspath(deleted_text_before_table_docx)
    
    # テーブルの内容を編集
    result = await edit_docx(abs_path, [
        {"paragraph_index": 1, "search": "Test1 | Test2", "replace": "Modified1 | Modified2"}
    ])
    
    # 変更を確認
    assert "-Test1 | Test2" in result
    assert "+Modified1 | Modified2" in result
    
    # ドキュメント構造が維持されていることを確認
    content = await read_docx(abs_path)
    assert "Modified1" in content
    assert "Modified2" in content
    # 削除されたテキストが表示されないことを確認
    assert "This text is deleted" not in content

def test_extract_table_text():
    """Test table text extraction."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "X1"
    table.cell(0, 1).text = "Y1"
    table.cell(1, 0).text = "X2"
    table.cell(1, 1).text = "Y2"
    
    result = extract_table_text(table)
    assert "X1 | Y1" in result
    assert "X2 | Y2" in result

@pytest.mark.asyncio
async def test_edit_docx_insert_basic(sample_docx):
    """Test basic paragraph insertion."""
    abs_path = os.path.abspath(sample_docx)
    
    result = await edit_docx_insert(abs_path, [
        {"text": "Inserted Text", "paragraph_index": 1}
    ])
    
    # 変更を確認
    content = await read_docx(abs_path)
    assert "Inserted Text" in content
    assert "+Inserted Text" in result

@pytest.mark.asyncio
async def test_edit_docx_insert_multiple(sample_docx):
    """Test inserting multiple paragraphs at different positions."""
    abs_path = os.path.abspath(sample_docx)
    
    result = await edit_docx_insert(abs_path, [
        {"text": "First Insert", "paragraph_index": 0},
        {"text": "Second Insert", "paragraph_index": 2}
    ])
    
    content = await read_docx(abs_path)
    assert "First Insert" in content
    assert "Second Insert" in content
    assert content.index("First Insert") < content.index("Second Insert")

@pytest.mark.asyncio
async def test_edit_docx_insert_same_position(sample_docx):
    """Test inserting multiple paragraphs at the same position."""
    abs_path = os.path.abspath(sample_docx)
    
    result = await edit_docx_insert(abs_path, [
        {"text": "First Same Position"},
        {"text": "Second Same Position"},
        {"text": "Third Same Position"}
    ])
    
    content = await read_docx(abs_path)
    # 指定順序で挿入されていることを確認
    assert content.index("First Same Position") < content.index("Second Same Position")
    assert content.index("Second Same Position") < content.index("Third Same Position")

@pytest.mark.asyncio
async def test_edit_docx_insert_at_end(sample_docx):
    """Test inserting paragraph at the end of document."""
    abs_path = os.path.abspath(sample_docx)
    
    result = await edit_docx_insert(abs_path, [
        {"text": "End of Document"}
    ])
    
    content = await read_docx(abs_path)
    assert "End of Document" in content
    assert content.rindex("End of Document") > content.rindex("Goodbye World")

@pytest.mark.asyncio
async def test_edit_docx_insert_before_table(complex_docx):
    """Test inserting paragraph before table."""
    abs_path = os.path.abspath(complex_docx)
    
    result = await edit_docx_insert(abs_path, [
        {"text": "Before Table Text", "paragraph_index": 4}
    ])
    
    content = await read_docx(abs_path)
    assert "Before Table Text" in content
    # テーブルの前に挿入されていることを確認
    table_index = content.index("[Table]")
    insert_index = content.index("Before Table Text")
    assert insert_index < table_index

@pytest.mark.asyncio
async def test_edit_docx_insert_errors(sample_docx):
    """Test error cases for paragraph insertion."""
    # 存在しないファイル
    with pytest.raises(ValueError) as exc_info:
        await edit_docx_insert(os.path.abspath("nonexistent.docx"), [{"text": "Test"}])
    assert "File not found" in str(exc_info.value)
    
    # 範囲外のインデックス
    abs_path = os.path.abspath(sample_docx)
    with pytest.raises(ValueError) as exc_info:
        await edit_docx_insert(abs_path, [{"text": "Test", "paragraph_index": 999}])
    assert "Paragraph index out of range" in str(exc_info.value)
