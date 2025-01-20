import os
from typing import Dict
from docx import Document
from docx.shared import Inches
from docx.table import Table
from mcp.server.lowlevel import Server, NotificationOptions
from mcp.server.stdio import stdio_server
from mcp.server.models import InitializationOptions
from mcp import types

server = Server("office-server")

async def validate_path(path: str) -> bool:
    if not os.path.isfile(path):
        return False
    elif path.endswith(".docx"):
        return True
    else:
        return False

def extract_table_text(table: Table) -> str:
    """Extract text from table with formatting."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)

async def read_docx(path: str) -> str:
    """Read docx file as text including tables.
    
    Args:
        path: relative path to target docx file
    Returns:
        str: Text representation of the document including tables
    """
    if not await validate_path(path):
        raise ValueError(f"Not a docx file: {path}")
    
    document = Document(path)
    content = []

    # Process paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            content.append(paragraph.text)
    
    # Process tables
    for table in document.tables:
        table_text = extract_table_text(table)
        content.append(f"[Table]\n{table_text}")
    
    # Process inline shapes (images)
    for paragraph in document.paragraphs:
        if paragraph._element.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            content.append("[Image]")

    return "\n\n".join(content)

async def write_docx(path: str, content: str) -> None:
    """Create a new docx file with the given content.
    
    Args:
        path: target path to create docx file
        content: text content to write
    """
    document = Document()
    
    # Split content into sections
    sections = content.split("\n\n")
    
    for section in sections:
        if section.startswith("[Table]"):
            # Create table from text representation
            table_content = section[7:].strip()  # Remove [Table] prefix
            rows = table_content.split("\n")
            if rows:
                num_columns = len(rows[0].split(" | "))
                table = document.add_table(rows=len(rows), cols=num_columns)
                
                for i, row in enumerate(rows):
                    cells = row.split(" | ")
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.strip()
        elif section.startswith("[Image]"):
            document.add_paragraph("[Image placeholder]")
        else:
            document.add_paragraph(section)
    
    document.save(path)

async def edit_docx(path: str, edits: list[Dict[str, str]]) -> Dict[str, str]:
    """Edit docx file by replacing multiple text occurrences.
    
    Args:
        path: path to target docx file
        edits: list of dictionaries containing 'search' and 'replace' pairs
            [{'search': 'text to find', 'replace': 'text to replace with'}, ...]
    Returns:
        dict: Containing original and modified text
    """
    if not await validate_path(path):
        raise ValueError(f"Not a docx file: {path}")
    
    # Read original content
    original = await read_docx(path)
    modified = original
    
    # Apply all edits sequentially
    not_found = []
    for edit in edits:
        search = edit['search']
        replace = edit['replace']
        if search not in modified:
            not_found.append(search)
            continue
        modified = modified.replace(search, replace)
    
    if not_found:
        raise ValueError(f"Search text not found: {', '.join(not_found)}")
    
    # Write modified content to file
    await write_docx(path, modified)
    
    return {
        "original": original,
        "modified": modified
    }

@server.list_tools()
async def list_tools() -> list[types.Tool]:
    return [
        types.Tool(
            name="read_docx",
            description=(
                "Read complete contents of a docx file including tables and images."
                "Use this tool when you want to read file endswith '.docx'."
                "Paragraphs are separated with two line breaks."
                "This tool convert images into placeholder [Image]."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string"}
                },
                "required": ["path"]
            }
        ),
        types.Tool(
            name="write_docx",
            description=(
                "Create a new docx file with given content."
                "Two line breaks in content represent new paragraph."
                "Table should starts with [Table], and separated with '|'."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "content": {"type": "string"}
                },
                "required": ["path", "content"]
            }
        ),
        types.Tool(
            name="edit_docx",
            description=(
                "Make multiple text replacements in a docx file. Accepts a list of search/replace pairs "
                "and applies them sequentially. Returns a git-style diff showing the changes made. "
                "Only works within allowed directories."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "edits": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "search": {"type": "string"},
                                "replace": {"type": "string"}
                            },
                            "required": ["search", "replace"]
                        }
                    }
                },
                "required": ["path", "edits"]
            }
        )
    ]

@server.call_tool()
async def call_tool(
    name: str,
    arguments: dict
) -> list[types.TextContent]:
    if name == "read_docx":
        content = await read_docx(arguments["path"])
        return [types.TextContent(type="text", text=content)]
    elif name == "write_docx":
        await write_docx(arguments["path"], arguments["content"])
        return [types.TextContent(type="text", text="Document created successfully")]
    elif name == "edit_docx":
        result = await edit_docx(arguments["path"], arguments["edits"])
        return [types.TextContent(type="text", text=f"Document edited successfully\nDiff:\n{result['original']}\n--->\n{result['modified']}")]
    raise ValueError(f"Tool not found: {name}")

async def run():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="office-file-server",
                server_version="0.1.0",
                capabilities=server.get_capabilities(
                    notification_options=NotificationOptions(),
                    experimental_capabilities={},
                ),
            )
        )

if __name__ == "__main__":
    import asyncio
    asyncio.run(run())
