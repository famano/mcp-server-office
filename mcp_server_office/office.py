import os
from typing import Dict
from docx import Document
from docx.table import Table
from docx.oxml import OxmlElement
from mcp.server.lowlevel import Server, NotificationOptions
from mcp.server.stdio import stdio_server
from mcp.server.models import InitializationOptions
from mcp import types
import difflib
from mcp_server_office.tools import READ_DOCX, WRITE_DOCX, EDIT_DOCX_PARAGRAPH

# WordML namespace constants
WORDML_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W_P = f"{{{WORDML_NS['w']}}}p"    # paragraph
W_TBL = f"{{{WORDML_NS['w']}}}tbl"  # table
W_R = f"{{{WORDML_NS['w']}}}r"    # run
W_T = f"{{{WORDML_NS['w']}}}t"    # text
W_DRAWING = f"{{{WORDML_NS['w']}}}drawing"  # drawing

server = Server("office-server")

async def validate_path(path: str) -> bool:
    if not os.path.isabs(path):
        raise ValueError(f"Not a absolute path: {path}")
    if not os.path.isfile(path):
        raise ValueError(f"File not found: {path}")
    elif path.endswith(".docx"):
        return True
    else:
        return False

def extract_table_text(table: Table) -> str:
    """Extract text from table with formatting."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_p_texts = [process_track_changes(paragraph._element).strip() for paragraph in cell.paragraphs]
            celltext = "<br>".join(cell_p_texts) #複数行に渡る場合、<br>で表現
            cells.append(celltext)
        rows.append(" | ".join(cells))
    return "\n".join(rows)

def create_table_from_text(text: str, props :any =None) -> Table:
    """add table from text representation. if props are passed, apply it to all cells"""
    rows = text.split("\n")
    temp_doc = Document()
    if rows:
        num_columns = len(rows[0].split(" | "))
        table = temp_doc.add_table(rows=len(rows), cols=num_columns)
                
    for i, row in enumerate(rows):
        cells = row.split(" | ")
        for j, cell in enumerate(cells):
            table.cell(i, j).text = ""
            new_run = table.cell(i, j).paragraphs[0].add_run(cell.strip()) #<br>が入って改行されている場合でも文字として処理してしまう。要検討。
            if props is not None:
                new_run._element.append(props)
    return table

def process_track_changes(element: OxmlElement) -> str:
    """Process track changes in a paragraph element."""
    text = ""
    for child in element:
        if child.tag == W_R:  # Normal run
            for run_child in child:
                if run_child.tag == W_T:
                    text += run_child.text if run_child.text else ""
        elif child.tag.endswith('ins'):  # Insertion
            inserted_text = ""
            for run in child.findall('.//w:t', WORDML_NS):
                inserted_text += run.text if run.text else ""
            if inserted_text:
                text += inserted_text
    return text

async def read_docx(path: str) -> str:
    """Read docx file as text including tables.
    
    Args:
        path: relative path to target docx file
    Returns:
        str: Text representation of the document including tables
    """
    if not await validate_path(path):
        raise ValueError(f"Not a docx file: {path}")
    
    document = Document(path)
    content = []

    paragraph_index = 0
    table_index = 0
    
    # 全要素を順番に処理
    for element in document._body._body:
        # パラグラフの処理
        if element.tag == W_P:
            paragraph = document.paragraphs[paragraph_index]
            paragraph_index += 1
            # 画像のチェック
            if paragraph._element.findall(f'.//{W_DRAWING}', WORDML_NS):
                content.append("[Image]")
            # テキストのチェック
            else:
                text = process_track_changes(paragraph._element)
                if text.strip():
                    content.append(text)
                else:
                    # 空行を抜くと編集時に困るので、空行でも追加
                    content.append("")
        # テーブルの処理
        elif element.tag == W_TBL:
            table = document.tables[table_index]
            table_index += 1
            table_text = extract_table_text(table)
            content.append(f"[Table]\n{table_text}")

    separator = [f"--- Paragraph {i} ---" for i in range(len(content))]
    
    result = []
    for i, p in enumerate(content):
        result.append(separator[i])
        result.append(p)
        
    return "\n".join(result)

async def write_docx(path: str, content: str) -> None:
    """Create a new docx file with the given content.
    
    Args:
        path: target path to create docx file
        content: text content to write
    """
    document = Document()
    
    # Split content into sections
    sections = content.split("\n\n")
    
    for section in sections:
        if section.startswith("[Table]"):
            table = create_table_from_text(section[7:].strip()) # Remove [Table] prefix
            document.element.body.append(table._element)
        elif section.startswith("[Image]"):
            document.add_paragraph("[Image placeholder]")
        else:
            document.add_paragraph(section)
    
    document.save(path)

async def edit_docx_insert():
    pass

async def edit_docx_paragraph(path: str, edits: list[Dict[str, str | int]]) -> str:
    """Edit docx file by replacing text.
    
    Args:
        path: path to target docx file
        edits: list of dictionaries containing search and replace text, and paragraph_index
            [{'search': 'text to find', 'replace': 'text to replace with', 'paragraph_index': 0}, ...]
            paragraph_index: 0-based index of the paragraph to edit (required)
            search: text to find
            replace: text to replace with
    Returns:
        str: A git-style diff showing the changes made
    """
    if not await validate_path(path):
        raise ValueError(f"Not a docx file: {path}")
    
    doc = Document(path)
    original = await read_docx(path)
    not_found = []

    # パラグラフとテーブルを順番に収集
    elements = []
    paragraph_count = 0
    table_count = 0
    for element in doc._body._body:
        if element.tag == W_P:
            elements.append(('p', doc.paragraphs[paragraph_count]))
            paragraph_count += 1
        elif element.tag == W_TBL:
            elements.append(('tbl', doc.tables[table_count]))
            table_count += 1

    for edit in edits:
        search = edit["search"]
        replace = edit["replace"]
        
        if "paragraph_index" not in edit:
            raise ValueError("paragraph_index is required")
            
        paragraph_index = edit["paragraph_index"]
        if paragraph_index >= len(elements):
            raise ValueError(f"Paragraph index out of range: {paragraph_index}")
        
        element_type, element = elements[paragraph_index]
        
        if element_type == 'p':
            paragraph = element
            if search not in paragraph.text:
                not_found.append(f"'{search}' in paragraph {paragraph_index}")
                continue

            # Store original XML element and get first run's properties
            original_element = paragraph._element
            first_run_props = None
            runs = original_element.findall(f'.//w:r', WORDML_NS)
            if runs:
                first_run = runs[0]
                if hasattr(first_run, 'rPr'):
                    first_run_props = first_run.rPr
            
            # Create new paragraph with the entire text
            new_paragraph = doc.add_paragraph()
            
            # Copy paragraph properties
            if original_element.pPr is not None:
                new_paragraph._p.append(original_element.pPr)
            
            # Replace text and create a single run with first run's properties
            new_text = process_track_changes(paragraph._element).replace(search, replace, 1)
            new_run = new_paragraph.add_run(new_text)
            if first_run_props is not None:
                new_run._element.append(first_run_props)
            
            # Replace original paragraph with new one
            original_element.getparent().replace(original_element, new_paragraph._element)
            
        elif element_type == 'tbl':
            # tableの場合、複数行に渡る書換では、特に行列が増減する場合、書式を保つことが困難なため、とりあえず0,0の書式を適用することとする。要検討。
            table = element
            table_paragraph = table._element.getprevious()
            table_text = extract_table_text(table)
            if search in table_text:
                # 既存tableを削除（親要素の参照を保持して安全に削除）
                parent = table._element.getparent()
                if parent is not None:
                    parent.remove(table._element)
                else:
                    # テーブルが文書のルート要素である場合（先頭の場合などにおそらく必要）
                    doc.element.body.remove(table._element)
                
                # Get first run's properties from the first cell
                first_run_props = None
                for paragraph in table.rows[0].cells[0].paragraphs:
                    for run in paragraph.runs:
                        if run._element.rPr is not None:
                            first_run_props = run._element.rPr
                            break
                
                new_text = table_text.replace(search, replace, 1)
                new_table = create_table_from_text(new_text, first_run_props)
                elements[paragraph_index] = ("tbl", new_table) # これがないと複数編集時に、あとの編集でtableがみつからなくなる
                if table_paragraph is not None:
                    table_paragraph.addnext(new_table._element)
                else:
                    # Noneの場合はtableの前がない、つまり先頭を意味する
                    doc.element.body.insert(0, new_table._element)
            else:
                not_found.append(f"'{search}' in table at paragraph {paragraph_index}")
            
    if not_found:
        raise ValueError(f"Search text not found: {', '.join(not_found)}")
    
    doc.save(path)
    
    # Read modified content and create diff
    modified = await read_docx(path)
    
    # 差分の生成
    diff_lines = []
    original_lines = [line for line in original.split("\n") if line.strip()]
    modified_lines = [line for line in modified.split("\n") if line.strip()]
    
    for line in difflib.unified_diff(original_lines, modified_lines, n=0):
        if line.startswith('---') or line.startswith('+++'):
            continue
        if line.startswith('-') or line.startswith('+'):
            diff_lines.append(line)
    
    return "\n".join(diff_lines) if diff_lines else ""

@server.list_tools()
async def list_tools() -> list[types.Tool]:
    return [READ_DOCX, EDIT_DOCX_PARAGRAPH, WRITE_DOCX]

@server.call_tool()
async def call_tool(
    name: str,
    arguments: dict
) -> list[types.TextContent]:
    if name == "read_docx":
        content = await read_docx(arguments["path"])
        return [types.TextContent(type="text", text=content)]
    elif name == "write_docx":
        await write_docx(arguments["path"], arguments["content"])
        return [types.TextContent(type="text", text="Document created successfully")]
    elif name == "edit_docx_paragraph":
        result = await edit_docx_paragraph(arguments["path"], arguments["edits"])
        return [types.TextContent(type="text", text=result)]
    raise ValueError(f"Tool not found: {name}")

async def run():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="office-file-server",
                server_version="0.1.0",
                capabilities=server.get_capabilities(
                    notification_options=NotificationOptions(),
                    experimental_capabilities={},
                ),
            )
        )

if __name__ == "__main__":
    import asyncio
    asyncio.run(run())
